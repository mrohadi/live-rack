#!/usr/bin/env python3
"""Generate manual test documentation for live-rack (Excel + Word)."""
import os
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = os.path.dirname(os.path.abspath(__file__))

# ---------------------------------------------------------------------------
# Test data. Each case: id, title, precondition, steps(list), expected, priority
# ---------------------------------------------------------------------------
MODULES = []

def mod(name, area, cases):
    MODULES.append({"name": name, "area": area, "cases": cases})

# 1. Onboarding / Self-service signup -------------------------------------
mod("1. Onboarding & Signup", "FE + BE", [
    ("TC-ONB-01", "Open signup page",
     "App running at http://localhost:5173",
     ["Navigate to /signup",
      "Observe page heading and form fields"],
     "AuthLayout shows title 'Create your workspace'. Fields: Company (required), Work email (required, type=email), Your name (optional). Button 'Create workspace'.",
     "High"),
    ("TC-ONB-02", "Submit signup with valid data",
     "On /signup",
     ["Enter Company 'Acme Co'",
      "Enter Work email 'owner@acme.test'",
      "Enter Your name 'Ada Lovelace'",
      "Click 'Create workspace'"],
     "POST /api/v1/signup returns 200 {org_id,user_id,status}. Success view shows 'Check your email' + 'Go to sign in' link. New tenant org + admin user provisioned in Zitadel.",
     "High"),
    ("TC-ONB-03", "Required field validation",
     "On /signup",
     ["Leave Company empty",
      "Click 'Create workspace'"],
     "Native HTML5 required validation blocks submit; no request sent.",
     "Medium"),
    ("TC-ONB-04", "Invalid email format",
     "On /signup",
     ["Enter email 'not-an-email'",
      "Submit"],
     "Email type validation blocks submit.",
     "Medium"),
    ("TC-ONB-05", "Duplicate company / email",
     "Org 'Acme Co' already exists",
     ["Submit signup reusing same company/email",
      "Observe error"],
     "Request fails (non-2xx). Inline alert: 'Signup failed. That email or company may already exist.'",
     "High"),
    ("TC-ONB-06", "Pending state on submit",
     "On /signup",
     ["Submit valid form",
      "Observe button during request"],
     "Button disabled, label 'Creating…' while isPending.",
     "Low"),
    ("TC-ONB-07", "Navigate to sign in from signup",
     "On /signup",
     ["Click 'Sign in' link"],
     "Routes to / (AuthGuard → Zitadel login).",
     "Low"),
])

# 2. Authentication --------------------------------------------------------
mod("2. Authentication & Session", "FE + BE (Zitadel OIDC)", [
    ("TC-AUTH-01", "Unauthenticated redirect to login",
     "No active session",
     ["Navigate to /",
      "Observe redirect"],
     "AuthGuard triggers Zitadel PKCE login redirect.",
     "High"),
    ("TC-AUTH-02", "Successful login + callback",
     "Valid Zitadel user admin@localhost / Admin123!",
     ["Complete Zitadel login",
      "Observe redirect to /callback then app"],
     "CallbackPage exchanges code, session established, lands on Dashboard. JWT carries project-roles claim.",
     "High"),
    ("TC-AUTH-03", "JIT org provisioning on first login",
     "First login of new Zitadel org user",
     ["Login as new-org user",
      "Call GET /api/v1/me"],
     "Gateway maps Zitadel org → org_id, JIT-provisions user row. /me returns role + permissions.",
     "High"),
    ("TC-AUTH-04", "Relative /ui/* redirect recovery",
     "Logged in",
     ["Navigate to /ui/console manually"],
     "Route redirects to / (Navigate replace).",
     "Low"),
    ("TC-AUTH-05", "Sign out",
     "Logged in, sidebar visible",
     ["Click sidebar Sign out button"],
     "auth.signoutRedirect() clears session, returns to Zitadel login.",
     "High"),
    ("TC-AUTH-06", "Expired/invalid token rejected by API",
     "Token expired or tampered",
     ["Send API request with stale JWT"],
     "Gateway returns 401; JWKS verification fails.",
     "High"),
    ("TC-AUTH-07", "MFA reflected in capabilities",
     "User with 2FA enrolled in Zitadel",
     ["Login with second factor",
      "Open Users page header"],
     "ID token amr shows otp/totp/webauthn; header shows '· 2FA on'.",
     "Medium"),
])

# 3. Dashboard -------------------------------------------------------------
mod("3. Dashboard / Overview", "FE + BE", [
    ("TC-DASH-01", "Load sales summary",
     "Logged in, seed data loaded",
     ["Navigate to / (Overview)"],
     "GET /api/v1/sales/summary. Three StatCards: Revenue (formatted $), Units sold, Orders.",
     "High"),
    ("TC-DASH-02", "Loading state",
     "Slow network",
     ["Open Overview before data resolves"],
     "Shows 'Loading sales…'.",
     "Low"),
    ("TC-DASH-03", "7-day revenue sparkline",
     "Sales data present",
     ["Observe Revenue · 7 days card"],
     "SVG polyline sparkline rendered with aria-label '7-day revenue sparkline'.",
     "Medium"),
    ("TC-DASH-04", "Revenue currency format",
     "revenue_cents present",
     ["Verify Revenue value"],
     "formatCents converts cents → currency string correctly.",
     "Medium"),
])

# 4. Map & Zones -----------------------------------------------------------
mod("4. Map & Zones", "FE + BE", [
    ("TC-MAP-01", "Load zones canvas",
     "Logged in, store has zones",
     ["Navigate to /map"],
     "GET /stores/:storeID/zones. Konva canvas renders zones. Header hint text shown.",
     "High"),
    ("TC-MAP-02", "Add a zone",
     "On /map",
     ["Click '+ Add zone'",
      "Type 'Front Display'",
      "Press Enter or click Add"],
     "POST /stores/:storeID/zones with defaults (200x120, color #6366f1, capacity 100). New zone appears; form resets.",
     "High"),
    ("TC-MAP-03", "Cancel add zone",
     "Add form open",
     ["Press Escape or click Cancel"],
     "Form closes, no zone created.",
     "Low"),
    ("TC-MAP-04", "Empty zone name blocked",
     "Add form open",
     ["Leave name blank, click Add"],
     "handleAddZone returns early; no POST.",
     "Medium"),
    ("TC-MAP-05", "Drag to move zone",
     "Zone exists",
     ["Drag a zone to new position"],
     "PUT /stores/:storeID/zones/:id with merged delta; position persists on reload.",
     "High"),
    ("TC-MAP-06", "Resize zone via corner",
     "Zone selected",
     ["Drag a corner handle"],
     "Width/height update via PUT.",
     "Medium"),
    ("TC-MAP-07", "Single select shows sidebar",
     "Zones present",
     ["Click one zone"],
     "ZoneDetailSidebar shows that zone's detail (GET /zones/:id).",
     "Medium"),
    ("TC-MAP-08", "Shift-click multi-select",
     "Zones present",
     ["Shift-click 2+ zones"],
     "Multiple selected; sidebar hidden (selectedZone null when >1).",
     "Low"),
    ("TC-MAP-09", "View tabs Zones/Heat/Items",
     "On /map",
     ["Click Heat then Items tabs"],
     "Canvas viewMode switches; active tab highlighted.",
     "Medium"),
    ("TC-MAP-10", "Live scan updates zone",
     "WS connected, scanner active",
     ["Record a 'place' scan into a zone"],
     "useScanStream patches zone: items+1 if valid+place, else misplaced+1; lastScan ts updates without reload.",
     "High"),
    ("TC-MAP-11", "Delete zone",
     "Zone exists (admin/manager)",
     ["Trigger DELETE /stores/:storeID/zones/:id"],
     "Zone removed; 200/204.",
     "Medium"),
])

# 5. Scanner ---------------------------------------------------------------
mod("5. Scanner PWA", "FE + BE", [
    ("TC-SCAN-01", "Open scanner page",
     "Logged in",
     ["Navigate to /scanner"],
     "Shows Scanner title, Start camera + Connect Zebra buttons, video element.",
     "High"),
    ("TC-SCAN-02", "Start/stop camera",
     "Camera permission granted",
     ["Click 'Start camera'",
      "Click 'Stop camera'"],
     "@zxing barcode reader attaches to video; toggles label. Camera stream stops on Stop.",
     "High"),
    ("TC-SCAN-03", "Camera barcode scan",
     "Camera active",
     ["Present a barcode to camera"],
     "onScan fires; SKU prepended to list (dedup consecutive); POST /scan sent.",
     "High"),
    ("TC-SCAN-04", "Zebra WebHID connect + scan",
     "Zebra USB scanner attached, Chrome",
     ["Click 'Connect Zebra'",
      "Select device in prompt",
      "Trigger a hardware scan"],
     "WebHID device opens; hardware scan decodes to SKU and posts.",
     "Medium"),
    ("TC-SCAN-05", "Scan list dedup + cap",
     "Camera active",
     ["Scan same SKU twice consecutively",
      "Scan >20 distinct SKUs"],
     "Consecutive duplicate not re-added; list capped at 20 (slice).",
     "Medium"),
    ("TC-SCAN-06", "Offline enqueue",
     "API unreachable / offline",
     ["Perform a scan while offline"],
     "POST fails → enqueueScan to IndexedDB; pending count increments.",
     "High"),
    ("TC-SCAN-07", "Flush queue on reconnect",
     "Pending offline scans exist",
     ["Restore connectivity (online event)"],
     "flushQueue replays queued scans to /scan; pendingCount returns to 0.",
     "High"),
    ("TC-SCAN-08", "Scan validation rules (BE)",
     "Zone with category rules",
     ["POST /stores/:storeID/scan/validate with mismatched category"],
     "ValidateResponse valid=false, code/reason populated (mis-scan).",
     "High"),
    ("TC-SCAN-09", "Dual-scan requirement (BE)",
     "Zone requires dual scan",
     ["POST validate without dual_scan_confirmed"],
     "requires_dual_scan=true; second confirmed scan passes.",
     "Medium"),
    ("TC-SCAN-10", "Valid place adjusts qty (BE)",
     "Valid scan, action=place",
     ["POST validate (valid place)"],
     "AdjustItemLocationQty applies delta; ScanRecorded event published to NATS.",
     "High"),
    ("TC-SCAN-11", "Scan rate limit",
     "Scanner sending bursts",
     ["Send >20 scan events/sec from one scanner"],
     "Gateway throttles at 20 scan events/s per scanner.",
     "Medium"),
])

# 6. Inventory -------------------------------------------------------------
mod("6. Inventory", "FE + BE", [
    ("TC-INV-01", "Load inventory table",
     "Logged in, seed data",
     ["Navigate to /inventory"],
     "GET /stores/:storeID/inventory. Table cols SKU, Name, Category, Status, Velocity, Qty.",
     "High"),
    ("TC-INV-02", "Filter by zone",
     "Inventory loaded",
     ["Select a zone in Zone filter"],
     "Rows filtered to that zone.",
     "Medium"),
    ("TC-INV-03", "Filter by status",
     "Inventory loaded",
     ["Select a status"],
     "Rows filtered by ITEM_STATUSES value.",
     "Medium"),
    ("TC-INV-04", "Filter by velocity band",
     "Inventory loaded",
     ["Select hot/warm/cold/dead"],
     "Rows filtered; velocity chip colors match band.",
     "Medium"),
    ("TC-INV-05", "Empty filter result",
     "Filters exclude all",
     ["Combine filters with no match"],
     "Shows 'No stock matches the current filters.'",
     "Low"),
    ("TC-INV-06", "Live qty update from scan",
     "WS connected",
     ["Record a scan affecting a listed SKU"],
     "useScanStream patches matching row qty live (qty-<sku> cell updates) without reload.",
     "High"),
    ("TC-INV-07", "Velocity classification",
     "Rows with varied scan rates",
     ["Inspect Velocity column"],
     "rowVelocity returns correct band per row metrics.",
     "Medium"),
])

# 7. Tasks -----------------------------------------------------------------
mod("7. Tasks (Kanban)", "FE + BE", [
    ("TC-TASK-01", "Load task board",
     "Logged in, tasks exist",
     ["Navigate to /tasks"],
     "GET /stores/:storeID/tasks. 4 columns (TASK_COLUMNS); count + hint in header.",
     "High"),
    ("TC-TASK-02", "Drag task to new column",
     "Task in column A",
     ["Drag card to column B (>6px)"],
     "PATCH /stores/:storeID/tasks/:id status updated; card moves; persists on reload.",
     "High"),
    ("TC-TASK-03", "Click vs drag threshold",
     "Task present",
     ["Click card without dragging"],
     "No move (6px activation constraint).",
     "Medium"),
    ("TC-TASK-04", "Drop on same column",
     "Task present",
     ["Drag and drop back to origin column"],
     "No mutation (status unchanged guard).",
     "Low"),
    ("TC-TASK-05", "Assign task",
     "Task present (manager/admin)",
     ["PATCH /stores/:storeID/tasks/:id/assignee"],
     "Assignee updated; reflected on TaskCard.",
     "Medium"),
    ("TC-TASK-06", "Empty/loading state",
     "No tasks or loading",
     ["Open tasks while loading"],
     "Shows 'Loading tasks…'; empty columns render.",
     "Low"),
])

# 8. Pipelines -------------------------------------------------------------
mod("8. Pipelines", "FE + BE", [
    ("TC-PIPE-01", "Empty state + create from template",
     "No pipelines",
     ["Navigate to /pipelines",
      "Click 'Start Item Restoration template'"],
     "POST /stores/:storeID/pipelines/from-template ('item-restoration'). Board appears.",
     "High"),
    ("TC-PIPE-02", "Load board, default first pipeline",
     "Pipelines exist",
     ["Navigate to /pipelines"],
     "GET pipelines list; auto-selects first; GET /pipelines/:id board (stages+cards).",
     "High"),
    ("TC-PIPE-03", "Switch pipeline",
     "Multiple pipelines",
     ["Choose another from selector"],
     "Board reloads for selected pipeline.",
     "Medium"),
    ("TC-PIPE-04", "Drag card across stage",
     "Board loaded",
     ["Drag card to another StageColumn"],
     "PATCH /pipelines/:id/cards/:cardID moves card to new stage_position.",
     "High"),
    ("TC-PIPE-05", "Same-stage drop guard",
     "Board loaded",
     ["Drop card back on its stage"],
     "No mutation.",
     "Low"),
    ("TC-PIPE-06", "Ageing / SLA badge",
     "Cards past SLA",
     ["Observe header + alerts"],
     "Header shows 'N ageing · past SLA'; ageingCount computed from cards.",
     "Medium"),
    ("TC-PIPE-07", "Bottleneck alert",
     "Board has bottleneck stage",
     ["Observe alert banner"],
     "Red alert with bottleneck stage name + ageing_count.",
     "Medium"),
])

# 9. Analytics -------------------------------------------------------------
mod("9. Analytics & Insights", "FE + BE (ClickHouse)", [
    ("TC-ANL-01", "Load analytics page",
     "Logged in, scan history",
     ["Navigate to /analytics"],
     "Signals/recommendations, Zone performance bars, 7×24 heatmap render.",
     "High"),
    ("TC-ANL-02", "Zone performance bars",
     "Scan data present",
     ["Observe Zone performance · scans"],
     "GET /analytics/zones; ZonePerfBars render per zone.",
     "Medium"),
    ("TC-ANL-03", "Activity heatmap",
     "Scan data present",
     ["Observe Activity heatmap · 7×24"],
     "GET /analytics/heatmap; 7×24 grid rendered.",
     "Medium"),
    ("TC-ANL-04", "Recommendations list",
     "Insight engine produced recs",
     ["Observe Signals · recommendations"],
     "SignalCards list recommendations with Apply action.",
     "High"),
    ("TC-ANL-05", "Apply a recommendation",
     "Recommendation present",
     ["Click Apply on a card"],
     "POST /recommendations/apply; applyingId disables that card during request; list refreshes.",
     "High"),
    ("TC-ANL-06", "Loading states",
     "Slow data",
     ["Open analytics pre-load"],
     "Shows 'Loading zones…' / 'Loading heatmap…'.",
     "Low"),
])

# 10. Integrations ---------------------------------------------------------
mod("10. Integrations & Webhooks", "FE + BE", [
    ("TC-INT-01", "Load integrations page",
     "Logged in (admin)",
     ["Navigate to /integrations"],
     "Marketplace/webhooks, Connectors, Webhook event log sections render.",
     "High"),
    ("TC-INT-02", "Connector catalog + webhook URLs",
     "On integrations",
     ["Inspect WebhookConfig rows"],
     "CONNECTOR_CATALOG listed with direction badges; inbound rows show webhook URL code.",
     "Medium"),
    ("TC-INT-03", "Toggle outbound push",
     "Outbound-capable connector",
     ["Check 'Push events' checkbox"],
     "onToggleOutbound updates state; disabled for comingSoon connectors.",
     "Medium"),
    ("TC-INT-04", "Connected connectors list",
     "Integration connected",
     ["Observe Connectors grid"],
     "GET /integrations; cards show kind, status tone, external_id. Empty → 'No integrations connected.'",
     "Medium"),
    ("TC-INT-05", "Inbound webhook delivery log",
     "Provider sent webhook",
     ["POST /webhooks/:provider (e.g. shopify)",
      "Refresh log"],
     "GET /integrations/webhooks; row shows provider, topic, event_id, status, received_at.",
     "High"),
    ("TC-INT-06", "Webhook idempotency",
     "Duplicate event_id",
     ["Re-POST same webhook event_id"],
     "Processed once; log not double-counted.",
     "Medium"),
    ("TC-INT-07", "Billing webhook",
     "Stripe billing event",
     ["POST /webhooks/billing"],
     "Handled by billing Webhook handler; 200.",
     "Medium"),
])

# 11. Users & Access -------------------------------------------------------
mod("11. Users & Access", "FE + BE", [
    ("TC-USR-01", "Load roster + capabilities",
     "Logged in",
     ["Navigate to /users"],
     "GET /users + GET /me. Member table + 'N members · you: <role>' header. Permission matrix shown.",
     "High"),
    ("TC-USR-02", "Invite button gated to admin",
     "Login as admin vs non-admin",
     ["Compare /users header for each role"],
     "canInvite true only when role==admin → 'Invite member' button shown; hidden for manager/staff/readonly.",
     "High"),
    ("TC-USR-03", "Invite a teammate",
     "Admin on /users",
     ["Click 'Invite member'",
      "Enter email, display name, pick role",
      "Click 'Send invite'"],
     "POST /api/v1/users/invite. Modal closes; roster refreshes; invited user appears/pending. Zitadel invite email sent.",
     "High"),
    ("TC-USR-04", "Invite role options",
     "Invite modal open",
     ["Open Role select"],
     "Options: admin, manager, staff, readonly (ASSIGNABLE_ROLES).",
     "Medium"),
    ("TC-USR-05", "Invite error handling",
     "Bad/duplicate email",
     ["Submit invalid invite"],
     "Inline alert 'Invite failed. Check the email and try again.'",
     "Medium"),
    ("TC-USR-06", "Resend invite",
     "Pending invite exists",
     ["POST /api/v1/users/:id/resend"],
     "Invite re-sent; 200.",
     "Low"),
    ("TC-USR-07", "Invite forbidden for non-admin (BE)",
     "Non-admin token",
     ["POST /users/invite as staff"],
     "API returns 403 (role gate, not per-request MFA).",
     "High"),
    ("TC-USR-08", "2FA indicator",
     "Admin with MFA",
     ["Observe header"],
     "Shows '· 2FA on' when amr has second factor.",
     "Low"),
    ("TC-USR-09", "Permission matrix accuracy",
     "On /users",
     ["Compare matrix vs role behavior"],
     "PERMISSION_MATRIX ticks match enforced capabilities (e.g. Edit users admin-only).",
     "Medium"),
])

# 12. Search ---------------------------------------------------------------
mod("12. Global Search (Command Palette)", "FE + BE", [
    ("TC-SRCH-01", "Open palette",
     "Logged in",
     ["Trigger command palette (search icon/shortcut)"],
     "Overlay opens; input focused; query reset.",
     "Medium"),
    ("TC-SRCH-02", "Search items + zones",
     "Palette open, seed data",
     ["Type query ≥ MIN_QUERY_LENGTH chars"],
     "GET /search; results listed with label, sublabel, kind (Item/Zone).",
     "High"),
    ("TC-SRCH-03", "Keyboard navigation",
     "Results present",
     ["Press ArrowDown/ArrowUp",
      "Press Enter"],
     "Active row moves within bounds; Enter selects and navigates (item→/inventory, zone→/map).",
     "Medium"),
    ("TC-SRCH-04", "Select by click",
     "Results present",
     ["Click a result"],
     "Palette closes; navigates to mapped route.",
     "Medium"),
    ("TC-SRCH-05", "Empty result",
     "Query has no match",
     ["Type non-matching query"],
     "Shows 'No matches.'",
     "Low"),
    ("TC-SRCH-06", "Escape closes",
     "Palette open",
     ["Press Escape or click overlay"],
     "Palette closes.",
     "Low"),
    ("TC-SRCH-07", "Min query length",
     "Palette open",
     ["Type 1 char"],
     "No fetch until MIN_QUERY_LENGTH reached.",
     "Low"),
])

# 13. Cross-cutting BE / Platform -----------------------------------------
mod("13. Platform / Cross-cutting (BE)", "BE", [
    ("TC-PLT-01", "Health check",
     "API running",
     ["GET /healthz"],
     "200 OK.",
     "High"),
    ("TC-PLT-02", "Prometheus metrics",
     "API running",
     ["GET /metrics"],
     "Prometheus exposition output.",
     "Low"),
    ("TC-PLT-03", "Swagger UI",
     "API running",
     ["GET /swagger/index.html"],
     "OpenAPI docs render.",
     "Low"),
    ("TC-PLT-04", "Multi-tenancy isolation (RLS)",
     "Two orgs A and B with data",
     ["Auth as org A",
      "Request org B resource ids"],
     "RLS prevents cross-org reads/writes; only org A rows returned.",
     "High"),
    ("TC-PLT-05", "Org-scoped repo args",
     "Any repo call",
     ["Inspect handler → store call"],
     "orgID passed explicitly; never derived from context alone.",
     "High"),
    ("TC-PLT-06", "API rate limit per org",
     "Burst traffic",
     ["Send >100 req/s for one org"],
     "Gateway throttles at 100 req/s per org (429).",
     "Medium"),
    ("TC-PLT-07", "Create service token",
     "Admin",
     ["POST /api/v1/service-tokens"],
     "Service-role token issued for machine access.",
     "Medium"),
    ("TC-PLT-08", "WebSocket gateway",
     "Logged in",
     ["Connect GET /api/v1/ws"],
     "WS upgrades; scan/recommendation events fan out to subscribed clients.",
     "High"),
    ("TC-PLT-09", "NATS → ClickHouse ingest",
     "Scan events flowing",
     ["Produce scan events",
      "Query ClickHouse"],
     "Ingest worker writes events to ClickHouse; analytics reflect them.",
     "Medium"),
    ("TC-PLT-10", "Rollup analytics job",
     "Historical events",
     ["Run rollup cron"],
     "Aggregates populate; heatmap/zone perf served from rollups.",
     "Low"),
])

# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------
def build_excel():
    wb = Workbook()
    head_fill = PatternFill("solid", fgColor="1F2937")
    head_font = Font(bold=True, color="FFFFFF", size=11)
    title_font = Font(bold=True, size=16, color="1F2937")
    thin = Side(style="thin", color="D1D5DB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    wrap = Alignment(wrap_text=True, vertical="top")
    center = Alignment(horizontal="center", vertical="center")
    prio_fill = {
        "High": PatternFill("solid", fgColor="FEE2E2"),
        "Medium": PatternFill("solid", fgColor="FEF3C7"),
        "Low": PatternFill("solid", fgColor="DCFCE7"),
    }

    # ---- Cover / Index sheet ----
    idx = wb.active
    idx.title = "Index"
    idx["A1"] = "live-rack — Manual Test Plan"
    idx["A1"].font = title_font
    idx["A2"] = "End-to-end manual test cases (Backend + Frontend). Generated for QA."
    idx["A2"].font = Font(italic=True, color="6B7280")
    headers = ["#", "Module", "Area", "Test Cases", "Sheet"]
    idx.append([])
    idx.append(headers)
    hrow = idx.max_row
    for c, h in enumerate(headers, 1):
        cell = idx.cell(row=hrow, column=c)
        cell.fill = head_fill; cell.font = head_font; cell.border = border; cell.alignment = center
    total = 0
    for i, m in enumerate(MODULES, 1):
        n = len(m["cases"]); total += n
        sheet = m["name"].split(".")[0].strip()  # number → sheet name
        sheet = f"M{sheet}"
        idx.append([i, m["name"], m["area"], n, sheet])
        for c in range(1, 6):
            idx.cell(row=idx.max_row, column=c).border = border
            idx.cell(row=idx.max_row, column=c).alignment = wrap
    idx.append(["", "TOTAL", "", total, ""])
    idx.cell(row=idx.max_row, column=2).font = Font(bold=True)
    idx.cell(row=idx.max_row, column=4).font = Font(bold=True)
    widths = [5, 38, 22, 12, 10]
    for c, w in enumerate(widths, 1):
        idx.column_dimensions[get_column_letter(c)].width = w
    idx.freeze_panes = "A5"

    # ---- Module sheets ----
    cols = ["Test ID", "Title", "Precondition", "Test Steps", "Expected Result", "Priority", "Status", "Notes"]
    cwidth = [12, 28, 30, 50, 50, 10, 12, 24]
    for m in MODULES:
        sheet = "M" + m["name"].split(".")[0].strip()
        ws = wb.create_sheet(title=sheet[:31])
        ws.cell(row=1, column=1, value=m["name"]).font = title_font
        ws.cell(row=2, column=1, value=f"Area: {m['area']}").font = Font(italic=True, color="6B7280")
        ws.append([])
        ws.append(cols)
        hrow = ws.max_row
        for c, h in enumerate(cols, 1):
            cell = ws.cell(row=hrow, column=c)
            cell.fill = head_fill; cell.font = head_font; cell.border = border; cell.alignment = center
        for tc in m["cases"]:
            tid, title, pre, steps, exp, prio = tc
            steps_str = "\n".join(f"{i+1}. {s}" for i, s in enumerate(steps))
            ws.append([tid, title, pre, steps_str, exp, prio, "Not Run", ""])
            r = ws.max_row
            for c in range(1, len(cols)+1):
                cell = ws.cell(row=r, column=c)
                cell.border = border; cell.alignment = wrap
            pcell = ws.cell(row=r, column=6)
            pcell.alignment = center
            if prio in prio_fill: pcell.fill = prio_fill[prio]
        for c, w in enumerate(cwidth, 1):
            ws.column_dimensions[get_column_letter(c)].width = w
        ws.freeze_panes = "A5"

    path = os.path.join(OUT, "live-rack-manual-test-plan.xlsx")
    wb.save(path)
    return path, total

# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------
def build_word(total):
    doc = Document()
    # base style
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    t = doc.add_heading("live-rack — Manual Test Plan", level=0)
    p = doc.add_paragraph("End-to-end manual test documentation (Backend + Frontend).")
    p.runs[0].italic = True
    doc.add_paragraph(f"Total modules: {len(MODULES)}   |   Total test cases: {total}")

    doc.add_heading("1. Scope & Environment", level=1)
    for line in [
        "Frontend: React 18 + Vite at http://localhost:5173",
        "Backend API: Go + Echo at http://localhost:8080",
        "Auth: Zitadel OIDC (console http://localhost:8081), admin@localhost / Admin123!",
        "Infra: PostgreSQL+TimescaleDB, NATS JetStream, ClickHouse, Redis, MinIO (make dev)",
        "Seed data: make seed",
        "Browser: Chrome (required for WebHID / Zebra scanner tests)",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("2. Preconditions (run once)", level=1)
    for i, step in enumerate([
        "Start infra: make dev",
        "Run migrations: goose -dir migrations postgres \"$DATABASE_URL\" up",
        "Seed fixtures: make seed",
        "Start API: (cd services/api && go run .)",
        "Start web: pnpm -F web dev",
        "Create a tenant via /signup or log in as admin@localhost.",
    ], 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    doc.add_heading("3. Priority & Status Legend", level=1)
    leg = doc.add_table(rows=1, cols=2); leg.style = "Light Grid Accent 1"
    leg.rows[0].cells[0].text = "Field"; leg.rows[0].cells[1].text = "Values"
    for k, v in [("Priority", "High / Medium / Low"),
                 ("Status", "Not Run / Pass / Fail / Blocked")]:
        row = leg.add_row().cells; row[0].text = k; row[1].text = v

    doc.add_heading("4. Test Modules", level=1)

    for m in MODULES:
        doc.add_heading(m["name"], level=2)
        sub = doc.add_paragraph(); r = sub.add_run(f"Area: {m['area']}"); r.italic = True; r.font.color.rgb = RGBColor(0x6B,0x72,0x80)
        for tc in m["cases"]:
            tid, title, pre, steps, exp, prio = tc
            h = doc.add_heading(f"{tid} — {title}", level=3)
            tbl = doc.add_table(rows=0, cols=2)
            tbl.style = "Light List Accent 1"
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            def addrow(label, value):
                cells = tbl.add_row().cells
                cells[0].text = label
                cells[0].paragraphs[0].runs[0].bold = True
                cells[1].text = value
            addrow("Priority", prio)
            addrow("Precondition", pre)
            steps_cell = tbl.add_row().cells
            steps_cell[0].text = "Steps"; steps_cell[0].paragraphs[0].runs[0].bold = True
            sc = steps_cell[1]
            sc.text = ""
            for i, s in enumerate(steps, 1):
                para = sc.paragraphs[0] if i == 1 else sc.add_paragraph()
                para.text = f"{i}. {s}"
            addrow("Expected Result", exp)
            addrow("Status", "Not Run")
            # set column widths
            for row in tbl.rows:
                row.cells[0].width = Inches(1.4)
                row.cells[1].width = Inches(5.4)
            doc.add_paragraph()

    path = os.path.join(OUT, "live-rack-manual-test-plan.docx")
    doc.save(path)
    return path

xlsx_path, total = build_excel()
docx_path = build_word(total)
print("XLSX:", xlsx_path)
print("DOCX:", docx_path)
print("TOTAL CASES:", total)
print("MODULES:", len(MODULES))
